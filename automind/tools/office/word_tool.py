"""Word 工具 —— 读写 .docx。

社区版动作：read / create / append / table / to_text
专业版动作（office_pro）：template（{{占位符}} 模板套打 / 邮件合并）
"""

from __future__ import annotations

from pathlib import Path
from typing import Any

from automind.core.types import PermissionTier, ToolResult
from automind.tools._toolkit import bad, delegate_pro, err, need, ok
from automind.tools.base import AbstractTool

#: 进阶动作（由专业版 office_pro 实现，社区版仅转交）
PRO_ACTIONS = {"template"}


class WordTool(AbstractTool):
    """读写 Word 文档（.docx）。"""

    name = "word_tool"
    description = (
        "Read and write Word documents (.docx). Actions: read (paragraphs + tables), "
        "create (new document from paragraphs/headings), append (add content), "
        "table (insert a table), to_text (plain-text export). "
        "The template action (mail-merge into {{placeholders}}) requires the Pro edition. "
        "Note: legacy .doc is not supported — convert to .docx first."
    )
    parameters = {
        "type": "object",
        "properties": {
            "action": {
                "type": "string",
                "enum": ["read", "create", "append", "table", "to_text", "template"],
            },
            "path": {"type": "string", "description": "Document path (.docx)."},
            "paragraphs": {
                "type": "array", "items": {"type": "string"},
                "description": "Paragraph texts for create/append.",
            },
            "heading": {"type": "string", "description": "Optional heading text."},
            "heading_level": {"type": "number", "description": "Heading level 1-9 (default 1)."},
            "rows": {
                "type": "array", "items": {"type": "array"},
                "description": "Table rows (first row treated as header) for the table action.",
            },
            "text_path": {"type": "string", "description": "Output path for to_text."},
            "max_paragraphs": {"type": "number", "description": "Cap for read (default 500)."},
        },
        "required": ["action", "path"],
    }
    permission_tier = PermissionTier.SENSITIVE
    risk_score = 30

    async def execute(self, **kwargs: Any) -> ToolResult:
        action = str(kwargs.get("action", "")).lower()
        path = Path(str(kwargs.get("path", ""))).expanduser()
        try:
            if action in PRO_ACTIONS:
                need("docx")
                return ok(self.name, **delegate_pro("office_pro", self.name, action, kwargs))
            need("docx")                      # 依赖检查（包名 python-docx，模块名 docx）
            import docx  # noqa: PLC0415 - 懒加载，见 _toolkit 说明

            if path.suffix.lower() == ".doc":
                return bad(self.name,
                           "不支持老式 .doc 二进制格式，请先另存为 .docx 再操作")
            if action == "create":
                return self._create(docx, path, kwargs)
            if action in ("append", "table"):
                return self._append(docx, path, kwargs, table_only=(action == "table"))
            if action == "read":
                return self._read(docx, path, kwargs)
            if action == "to_text":
                return self._to_text(docx, path, kwargs)
            return bad(self.name, f"不支持的 action：{action}")
        except Exception as e:
            return err(self.name, e)

    # ── 具体动作 ──────────────────────────────────────────

    @staticmethod
    def _add_content(doc: Any, kw: dict) -> int:
        n = 0
        if kw.get("heading"):
            doc.add_heading(str(kw["heading"]), level=int(kw.get("heading_level") or 1))
            n += 1
        for p in kw.get("paragraphs") or []:
            doc.add_paragraph(str(p))
            n += 1
        rows = kw.get("rows") or []
        if rows:
            table = doc.add_table(rows=len(rows), cols=max(len(r) for r in rows))
            table.style = "Table Grid"
            for i, row in enumerate(rows):
                for j, val in enumerate(row):
                    table.cell(i, j).text = "" if val is None else str(val)
            n += len(rows)
        return n

    def _create(self, docx: Any, path: Path, kw: dict) -> ToolResult:
        doc = docx.Document()
        n = self._add_content(doc, kw)
        path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(path)
        return ok(self.name, path=str(path), blocks=n,
                  message=f"已创建文档 {path.name}（{n} 个内容块）")

    def _append(self, docx: Any, path: Path, kw: dict, table_only: bool) -> ToolResult:
        if table_only and not kw.get("rows"):
            return bad(self.name, "table 动作需要提供 rows")
        doc = docx.Document(path) if path.is_file() else docx.Document()
        n = self._add_content(doc, kw)
        if n == 0:
            return bad(self.name, "没有可追加的内容（paragraphs / heading / rows 均为空）")
        path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(path)
        return ok(self.name, path=str(path), added=n,
                  message=f"已向 {path.name} 追加 {n} 个内容块")

    def _read(self, docx: Any, path: Path, kw: dict) -> ToolResult:
        if not path.is_file():
            return bad(self.name, f"文件不存在：{path}")
        cap = int(kw.get("max_paragraphs") or 500)
        doc = docx.Document(path)
        paras = [p.text for p in doc.paragraphs if p.text.strip()]
        truncated = len(paras) > cap
        tables = [[[c.text for c in row.cells] for row in t.rows] for t in doc.tables]
        return ok(self.name, path=str(path), paragraphs=paras[:cap],
                  paragraph_count=len(paras), truncated=truncated,
                  tables=tables, table_count=len(tables))

    def _to_text(self, docx: Any, path: Path, kw: dict) -> ToolResult:
        if not path.is_file():
            return bad(self.name, f"文件不存在：{path}")
        doc = docx.Document(path)
        lines = [p.text for p in doc.paragraphs]
        for t in doc.tables:
            for row in t.rows:
                lines.append("\t".join(c.text for c in row.cells))
        text = "\n".join(lines)
        out = kw.get("text_path")
        if out:
            p = Path(str(out)).expanduser()
            p.parent.mkdir(parents=True, exist_ok=True)
            p.write_text(text, encoding="utf-8")
            return ok(self.name, path=str(p), chars=len(text),
                      message=f"已导出纯文本到 {p.name}")
        return ok(self.name, text=text[:20000], chars=len(text),
                  truncated=len(text) > 20000)
