"""批量文档技能 —— 批量合并 / 提取文本 / 重命名 Word、PDF。

复用 pypdf / python-docx（可选依赖），对目录内同类文档做批处理：
- merge：把目录下所有 PDF（或所有 .docx）合并成一个文件；
- to_text：把每个 Word/PDF 提取为同名 .txt；
- rename：按前缀/后缀/序号批量重命名。
"""

from __future__ import annotations

from pathlib import Path
from typing import Any

from pydantic import BaseModel

from automind.skills.skill_base import AbstractSkill, SkillResult


class DocBatchInput(BaseModel):
    directory: str
    operation: str = "merge"  # merge | to_text | rename
    kind: str = "pdf"  # pdf | word（merge/to_text 的目标类型）
    output: str = ""  # merge 的输出文件
    prefix: str = ""  # rename 前缀
    suffix: str = ""  # rename 后缀
    pattern: str = "*"  # 文件匹配（默认全量）


class DocBatchSkill(AbstractSkill):
    """批量处理目录下的 Word / PDF 文档。"""

    name = "doc_batch"
    description = "Batch-process Word/PDF files in a directory: merge, extract text, rename"

    async def execute(self, input_data: Any, agent: Any = None) -> SkillResult:
        inp = DocBatchInput(**input_data) if isinstance(input_data, dict) else input_data
        root = Path(inp.directory).expanduser()
        try:
            if not root.is_dir():
                return SkillResult(success=False, error=f"目录不存在：{root}")

            if inp.operation == "merge":
                return await self._merge(root, inp)
            if inp.operation == "to_text":
                return await self._to_text(root, inp)
            if inp.operation == "rename":
                return self._rename(root, inp)
            return SkillResult(success=False, error=f"未知 operation：{inp.operation}")
        except Exception as e:
            return SkillResult(success=False, error=str(e))

    async def _merge(self, root: Path, inp: DocBatchInput) -> SkillResult:
        from automind.tools._toolkit import need
        suffix = ".docx" if inp.kind == "word" else ".pdf"
        files = sorted([p for p in root.glob(f"{inp.pattern}{suffix}")])
        if len(files) < 2:
            return SkillResult(success=False, error=f"找到的 {suffix} 文件不足 2 个：{len(files)}")
        out = Path(inp.output).expanduser() if inp.output else root / f"merged{suffix}"
        out.parent.mkdir(parents=True, exist_ok=True)

        if inp.kind == "word":
            need("docx")
            import docx
            merged = docx.Document()
            for i, f in enumerate(files):
                if i > 0:
                    merged.add_page_break()
                src = docx.Document(str(f))
                for para in src.paragraphs:
                    merged.add_paragraph(para.text)
            merged.save(str(out))
        else:
            need("pypdf")
            from pypdf import PdfWriter
            writer = PdfWriter()
            for f in files:
                writer.append(str(f))
            with out.open("wb") as fh:
                writer.write(fh)
        return SkillResult(success=True,
                           output=f"已合并 {len(files)} 个文件 → {out}",
                           artifacts=[str(out)])

    async def _to_text(self, root: Path, inp: DocBatchInput) -> SkillResult:
        from automind.tools._toolkit import need
        suffix = ".docx" if inp.kind == "word" else ".pdf"
        files = sorted([p for p in root.glob(f"{inp.pattern}{suffix}")])
        if not files:
            return SkillResult(success=False, error=f"未找到 {suffix} 文件")
        artifacts: list[str] = []
        for f in files:
            if inp.kind == "word":
                need("docx")
                import docx
                doc = docx.Document(str(f))
                text = "\n".join(p.text for p in doc.paragraphs)
            else:
                need("pypdf")
                from pypdf import PdfReader
                text = "\n".join((pg.extract_text() or "") for pg in PdfReader(str(f)).pages)
            txt = f.with_suffix(".txt")
            txt.write_text(text, encoding="utf-8")
            artifacts.append(str(txt))
        return SkillResult(success=True, output=f"已提取 {len(artifacts)} 个文本文件", artifacts=artifacts)

    @staticmethod
    def _rename(root: Path, inp: DocBatchInput) -> SkillResult:
        files = sorted([p for p in root.glob(inp.pattern) if p.is_file()])
        renamed: list[str] = []
        for i, f in enumerate(files, 1):
            target = root / f"{inp.prefix}{i:03d}{inp.suffix}{f.suffix}"
            f.rename(target)
            renamed.append(target.name)
        return SkillResult(success=True, output=f"已重命名 {len(renamed)} 个文件", artifacts=[str(root)])
